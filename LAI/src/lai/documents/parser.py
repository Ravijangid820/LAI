"""Document parser using Docling for PDF/DOCX.

Extracts text with structural preservation (headings, sections).
"""

from dataclasses import dataclass, field
from pathlib import Path

from lai.core.config import get_settings
from lai.core.exceptions import FileTooLargeError, UnsupportedFormatError
from lai.core.logging import get_logger

logger = get_logger("lai.documents.parser")

SUPPORTED_EXTENSIONS = {".pdf", ".docx"}


@dataclass
class ParsedDocument:
    title: str = ""
    text: str = ""
    sections: list[dict] = field(default_factory=list)
    page_count: int = 0
    metadata: dict = field(default_factory=dict)


def parse_document(file_path: str | Path) -> ParsedDocument:
    """Parse a PDF or DOCX file into structured text.

    Uses Docling for layout-aware parsing with heading detection.
    """
    path = Path(file_path)
    ext = path.suffix.lower()

    if ext not in SUPPORTED_EXTENSIONS:
        raise UnsupportedFormatError(f"Unsupported format: {ext}. Supported: {SUPPORTED_EXTENSIONS}")

    settings = get_settings().chunking
    file_size_mb = path.stat().st_size / (1024 * 1024)
    if file_size_mb > settings.max_file_size_mb:
        raise FileTooLargeError(f"File {path.name} is {file_size_mb:.1f}MB (max: {settings.max_file_size_mb}MB)")

    logger.info("Parsing document: %s (%.1fMB)", path.name, file_size_mb)

    try:
        from docling.document_converter import DocumentConverter

        converter = DocumentConverter()
        result = converter.convert(str(path))
        doc = result.document

        # Extract full text
        full_text = doc.export_to_markdown()

        # Extract sections from headings
        sections = []
        for item in doc.iterate_items():
            if hasattr(item, 'label') and 'heading' in str(item.label).lower():
                sections.append({"heading": item.text, "level": getattr(item, 'level', 1)})

        parsed = ParsedDocument(
            title=doc.name if hasattr(doc, 'name') else path.stem,
            text=full_text,
            sections=sections,
            page_count=getattr(doc, 'num_pages', 0),
            metadata={"source": str(path), "format": ext},
        )

        logger.info("Parsed %s: %d chars, %d sections, %d pages", path.name, len(full_text), len(sections), parsed.page_count)
        return parsed

    except ImportError:
        logger.warning("Docling not available, falling back to basic text extraction")
        return _basic_parse(path)
    except Exception as e:
        logger.error("Docling parsing failed for %s: %s", path.name, e)
        return _basic_parse(path)


def _basic_parse(path: Path) -> ParsedDocument:
    """Fallback parser when Docling is unavailable."""
    ext = path.suffix.lower()

    if ext == ".docx":
        try:
            from docx import Document as DocxDocument

            doc = DocxDocument(str(path))
            text = "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
            return ParsedDocument(title=path.stem, text=text, metadata={"source": str(path), "format": ext})
        except Exception as e:
            logger.error("DOCX fallback failed: %s", e)

    # Last resort: try reading as text
    try:
        text = path.read_text(encoding="utf-8", errors="replace")
        return ParsedDocument(title=path.stem, text=text, metadata={"source": str(path), "format": ext})
    except Exception as e:
        logger.error("Text fallback failed: %s", e)
        return ParsedDocument(title=path.stem, text="", metadata={"error": str(e)})
