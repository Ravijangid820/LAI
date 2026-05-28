"""Render a stored DDiQ report into a Word (.docx) deliverable.

The report row's ``report_data`` JSONB (shape: :class:`ddiq.models.DDiQReportData`)
is the input. We work off the plain dict rather than the Pydantic model because
that is exactly what psycopg2 hands back from the JSONB column, and because the
report is populated incrementally — any field may be missing or ``None`` on a
report whose pipeline crashed midway, and the writer must never raise on that.

Output is a German, lawyer-facing document: a Kanzlei-Briefkopf placeholder in
the page header (firm drops in their letterhead), the project Eckdaten, findings
grouped by traffic-light severity with evidence and legal basis, the per-park
breakdown when a data room covers several sites, deadlines, Grundbuch and
Rückbau checks, and the list of analysed documents. It closes with the
AI-generated-output disclaimer a lawyer needs before forwarding it to a client.
"""

from __future__ import annotations

from io import BytesIO
from typing import Any, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

# Traffic-light → (German label, text colour, bullet glyph). The colours are
# print-legible (the yellow is a dark goldenrod, not a pale amber that vanishes
# on paper).
_AMPEL = {
    "red": ("Kritisch", RGBColor(0xC0, 0x00, 0x00), "●"),
    "yellow": ("Erhöhte Aufmerksamkeit", RGBColor(0xB8, 0x86, 0x0B), "●"),
    "green": ("Unkritisch", RGBColor(0x1E, 0x7D, 0x32), "●"),
}
_AMPEL_ORDER = ["red", "yellow", "green"]

_GREY = RGBColor(0x80, 0x80, 0x80)


# ── small value helpers ──────────────────────────────────────────────


def _s(value: Any) -> str:
    """Display string for a possibly-missing scalar."""
    if value is None:
        return ""
    return str(value).strip()


def _bundesland(value: Any) -> str:
    v = _s(value)
    return v.title() if v else ""


def _eur(value: Any) -> str:
    try:
        # German thousands separator (1.234.567 €).
        return f"{float(value):,.0f} €".replace(",", ".")
    except (TypeError, ValueError):
        return ""


def _mw(value: Any) -> str:
    try:
        return f"{float(value):.1f} MW"
    except (TypeError, ValueError):
        return ""


# ── paragraph / run helpers ──────────────────────────────────────────


def _label_value(doc: Document, label: str, value: str, *, color: Optional[RGBColor] = None) -> None:
    """A bold label followed by a plain value on one paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(f"{label}: ")
    run.bold = True
    val = p.add_run(value or "—")
    if color is not None:
        val.font.color.rgb = color


def _facts_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    """Two-column key/value table; skips rows whose value is empty."""
    rows = [(k, v) for k, v in rows if v]
    if not rows:
        doc.add_paragraph("Keine Eckdaten aus den Dokumenten bestimmbar.")
        return
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    for i, (k, v) in enumerate(rows):
        c0, c1 = table.rows[i].cells
        kp = c0.paragraphs[0]
        kr = kp.add_run(k)
        kr.bold = True
        c1.paragraphs[0].add_run(v)


def _evidence_lines(doc: Document, evidence: list[dict]) -> None:
    for ev in evidence or []:
        if not isinstance(ev, dict):
            continue
        bits = []
        clause = _s(ev.get("clause"))
        fname = _s(ev.get("doc_filename"))
        page = ev.get("page")
        if clause:
            bits.append(clause)
        if fname:
            bits.append(fname + (f", S. {page}" if page else ""))
        head = " — ".join(bits) if bits else "Beleg"
        excerpt = _s(ev.get("excerpt"))
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(f"Beleg: {head}")
        r.italic = True
        r.font.size = Pt(9)
        r.font.color.rgb = _GREY
        if excerpt:
            er = p.add_run(f" „{excerpt}“")
            er.italic = True
            er.font.size = Pt(9)
            er.font.color.rgb = _GREY


# ── sections ─────────────────────────────────────────────────────────


def _add_letterhead(doc: Document) -> None:
    """Firm letterhead placeholder in the page header (repeats every page)."""
    header = doc.sections[0].header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("[ KANZLEI-BRIEFKOPF — Logo und Anschrift der Kanzlei hier einfügen ]")
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = _GREY


def _add_title(doc: Document, report: dict, project_name: Optional[str]) -> None:
    name = project_name or _s(report.get("projectName")) or "Unbenanntes Projekt"
    title = doc.add_heading("Due-Diligence-Bericht", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    sub = doc.add_heading(name, level=1)
    sub.alignment = WD_ALIGN_PARAGRAPH.LEFT

    _facts_table(
        doc,
        [
            ("Projekt", name),
            ("Erstellt für", _s(report.get("preparedFor"))),
            ("Erstellt von", _s(report.get("preparedBy"))),
            ("Datum", _s(report.get("date"))),
        ],
    )
    doc.add_paragraph()


def _add_eckdaten(doc: Document, report: dict) -> None:
    facts = report.get("projectFacts") or {}
    doc.add_heading("1. Eckdaten des Projekts", level=1)
    _facts_table(
        doc,
        [
            ("Projektgesellschaft", _s(facts.get("projectCompany"))),
            ("Bundesland", _bundesland(facts.get("bundesland") or report.get("bundesland"))),
            ("Anzahl WEA", _s(facts.get("turbineCount") or report.get("turbineCount") or "")),
            ("Davon errichtet/abgenommen", _s(facts.get("commissionedWeaCount") or "")),
            ("Gesamtleistung", _mw(facts.get("totalCapacityMw"))),
        ],
    )
    # Honest "unknown" notes carried per field (Path B).
    notes = facts.get("notes") or {}
    if isinstance(notes, dict) and notes:
        doc.add_paragraph()
        np = doc.add_paragraph()
        np.add_run("Hinweise zu nicht bestimmbaren Angaben:").bold = True
        for field, note in notes.items():
            if _s(note):
                doc.add_paragraph(f"{field}: {note}", style="List Bullet")


def _add_findings(doc: Document, findings: list[dict], heading: str) -> None:
    findings = [f for f in (findings or []) if isinstance(f, dict)]
    if not findings:
        return
    doc.add_heading(heading, level=1)

    multi_park = any(_s(f.get("park")) for f in findings)
    buckets: dict[str, list[dict]] = {k: [] for k in _AMPEL_ORDER}
    other: list[dict] = []
    for f in findings:
        sev = _s(f.get("severity") or f.get("ampel")).lower()
        (buckets[sev] if sev in buckets else other).append(f)

    for sev in _AMPEL_ORDER:
        bucket = buckets[sev]
        if not bucket:
            continue
        label, color, glyph = _AMPEL[sev]
        sh = doc.add_heading(level=2)
        gr = sh.add_run(f"{glyph} {label} ({len(bucket)})")
        gr.font.color.rgb = color
        for f in bucket:
            _add_finding_block(doc, f, color, multi_park)
    for f in other:
        _add_finding_block(doc, f, None, multi_park)


def _add_finding_block(doc: Document, f: dict, color: Optional[RGBColor], multi_park: bool) -> None:
    domain = _s(f.get("domain")) or "Allgemein"
    head = doc.add_paragraph()
    hr = head.add_run(domain)
    hr.bold = True
    if color is not None:
        hr.font.color.rgb = color
    # Park attribution: a finding about a neighbouring park must not read as a
    # fact about the report's subject.
    park = _s(f.get("park"))
    if multi_park and park:
        pr = head.add_run(f"  (betrifft {park})")
        pr.italic = True
        pr.font.size = Pt(9)
        pr.font.color.rgb = _GREY

    doc.add_paragraph(_s(f.get("text")))

    if _s(f.get("legal_basis")):
        _label_value(doc, "Rechtsgrundlage", _s(f.get("legal_basis")))
    if _s(f.get("recommended_action")):
        _label_value(doc, "Empfohlene Maßnahme", _s(f.get("recommended_action")))

    q = f.get("quantification") or {}
    if isinstance(q, dict):
        qbits = []
        if q.get("mw_affected") is not None:
            qbits.append(f"betroffene Leistung {_mw(q.get('mw_affected'))}")
        if q.get("eur_impact_estimate") is not None:
            qbits.append(f"geschätzter Effekt {_eur(q.get('eur_impact_estimate'))}")
        if q.get("days_until_deadline") is not None:
            qbits.append(f"Frist in {q.get('days_until_deadline')} Tagen")
        if qbits:
            _label_value(doc, "Quantifizierung", "; ".join(qbits))

    _evidence_lines(doc, f.get("evidence") or [])
    doc.add_paragraph()


def _add_parks(doc: Document, report: dict) -> None:
    if not report.get("multiParkDetected"):
        return
    parks = [p for p in (report.get("parks") or []) if isinstance(p, dict)]
    if not parks:
        return
    doc.add_heading("Windparks im Datenraum", level=1)
    doc.add_paragraph(
        "Der Datenraum umfasst mehrere Windparks. Gegenstand dieses Berichts ist "
        "der als „Subjekt“ markierte Park; die übrigen werden nachrichtlich "
        "aufgeführt."
    )
    for p in parks:
        name = _s(p.get("name")) or "Unbenannter Park"
        h = doc.add_heading(level=2)
        hr = h.add_run(name + ("  (Subjekt)" if p.get("isPrimary") else ""))
        if not p.get("isPrimary"):
            hr.font.color.rgb = _GREY
        _facts_table(
            doc,
            [
                ("Projektgesellschaft", _s(p.get("projectCompany"))),
                ("Bundesland", _bundesland(p.get("bundesland"))),
                ("Standort", _s(p.get("location"))),
                ("Anzahl WEA", _s(p.get("turbineCount") or "")),
                ("Gesamtleistung", _mw(p.get("totalCapacityMw"))),
                ("Anlagentypen", ", ".join(p.get("models") or [])),
            ],
        )


def _add_timeline(doc: Document, report: dict) -> None:
    entries = [t for t in (report.get("timeline") or []) if isinstance(t, dict)]
    if not entries:
        return
    doc.add_heading("Fristen und Termine", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for cell, text in zip(hdr, ("Datum", "Vorgang", "Dringlichkeit")):
        cell.paragraphs[0].add_run(text).bold = True
    for t in entries:
        cells = table.add_row().cells
        cells[0].paragraphs[0].add_run(_s(t.get("date")))
        desc = _s(t.get("description"))
        if _s(t.get("legal_basis")):
            desc = f"{desc} ({_s(t.get('legal_basis'))})"
        cells[1].paragraphs[0].add_run(desc)
        cells[2].paragraphs[0].add_run(_s(t.get("urgency")))


def _add_grundbuch(doc: Document, report: dict) -> None:
    checks = [g for g in (report.get("grundbuchChecks") or []) if isinstance(g, dict)]
    if not checks:
        return
    doc.add_heading("Grundbuch-Abgleich", level=1)
    for g in checks:
        match = g.get("owner_match")
        match_txt = {True: "Eigentümer = Verpächter", False: "Abweichung Eigentümer/Verpächter"}.get(
            match, "nicht bestimmbar"
        )
        color = {True: _AMPEL["green"][1], False: _AMPEL["red"][1]}.get(match)
        _label_value(doc, _s(g.get("parcel_id")) or "Flurstück", match_txt, color=color)
        if _s(g.get("registered_owner")):
            doc.add_paragraph(f"Eingetragener Eigentümer: {_s(g.get('registered_owner'))}", style="List Bullet")
        if _s(g.get("lessor_name")):
            doc.add_paragraph(f"Verpächter laut Vertrag: {_s(g.get('lessor_name'))}", style="List Bullet")
        for enc in g.get("encumbrances") or []:
            doc.add_paragraph(f"Belastung: {_s(enc)}", style="List Bullet")
        _evidence_lines(doc, g.get("evidence") or [])


def _add_rueckbau(doc: Document, report: dict) -> None:
    bond = report.get("rueckbauBond")
    if not isinstance(bond, dict) or not bond:
        return
    doc.add_heading("Rückbaubürgschaft (§ 35 Abs. 5 BauGB)", level=1)
    suff = bond.get("sufficient")
    color = {True: _AMPEL["green"][1], False: _AMPEL["red"][1]}.get(suff)
    _facts_table(
        doc,
        [
            ("Betrag", _eur(bond.get("amount_eur"))),
            ("Art", _s(bond.get("instrument_type"))),
            ("Sicherungsgeber", _s(bond.get("provider"))),
            ("Begünstigte", _s(bond.get("beneficiary"))),
            ("Gültig bis", _s(bond.get("valid_until"))),
        ],
    )
    if suff is not None:
        _label_value(
            doc,
            "Bewertung",
            "voraussichtlich ausreichend" if suff else "voraussichtlich nicht ausreichend",
            color=color,
        )
    if _s(bond.get("note")):
        doc.add_paragraph(_s(bond.get("note")))


def _add_jurisdiction(doc: Document, report: dict) -> None:
    warnings = [w for w in (report.get("jurisdictionWarnings") or []) if isinstance(w, dict)]
    if not warnings:
        return
    doc.add_heading("Hinweise zur Rechtslage (Bundesland)", level=1)
    for w in warnings:
        text = _s(w.get("message") or w.get("rule") or w.get("text"))
        if text:
            doc.add_paragraph(text, style="List Bullet")


def _add_documents(doc: Document, report: dict) -> None:
    docs = report.get("analyzedDocuments") or []
    if not docs:
        dmap = report.get("documentMap") or []
        docs = [d.get("filename") for d in dmap if isinstance(d, dict) and d.get("filename")]
    if not docs:
        return
    doc.add_heading("Geprüfte Dokumente", level=1)
    for name in docs:
        if _s(name):
            doc.add_paragraph(_s(name), style="List Bullet")


def _add_disclaimer(doc: Document) -> None:
    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run(
        "Hinweis: Dieser Bericht wurde automatisiert auf Basis der bereitgestellten "
        "Dokumente erstellt und ersetzt keine rechtliche Beratung. Die Angaben sind "
        "vor Verwendung anhand der Originalquellen zu verifizieren."
    )
    r.italic = True
    r.font.size = Pt(8)
    r.font.color.rgb = _GREY


# ── public entry point ───────────────────────────────────────────────


def build_report_docx(report: dict, project_name: Optional[str] = None) -> bytes:
    """Render ``report`` (a DDiQReportData-shaped dict) to .docx bytes."""
    report = report or {}
    doc = Document()

    _add_letterhead(doc)
    _add_title(doc, report, project_name)
    _add_eckdaten(doc, report)
    _add_findings(doc, report.get("findings") or [], "2. Feststellungen")
    _add_findings(
        doc, report.get("crossDocFindings") or [], "3. Dokumentenübergreifende Feststellungen"
    )
    _add_parks(doc, report)
    _add_timeline(doc, report)
    _add_grundbuch(doc, report)
    _add_rueckbau(doc, report)
    _add_jurisdiction(doc, report)
    _add_documents(doc, report)
    _add_disclaimer(doc)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
